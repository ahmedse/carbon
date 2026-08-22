"""``export_document`` — generate Word (.docx) / Excel (.xlsx) artifacts.

This is the missing "deliverable" primitive for the workspace chat agent:
after a study, audit, or research task the agent can produce a real,
downloadable report instead of only chat prose.  ``format`` selects the
output (``docx``, ``xlsx``, or ``both``); content is accepted either as
markdown ``content`` (rendered to paragraphs / headings / bullets) and/or a
``table`` (headers + rows → a sheet / Word table).

Files are written under ``MEDIA_ROOT/ai_exports/`` and surfaced to chat as a
``download`` action — the UI renders a real download link (never a raw
server path the user must copy).

Guardrails honored (non-negotiable):

  * **RULE_20** — zero upward imports: only stdlib + ``django.conf`` (for
    ``MEDIA_ROOT``) + ``openpyxl`` / ``docx``.  No domain-app models/views.
  * **RULE_21** — file generation is **non-mutating to user data**; it writes
    a fresh artifact into a scratch media folder, so
    ``requires_confirmation=False`` (the user explicitly asked for the export;
    nothing in their records is created or changed).
  * **RULE_23** — outcome copy: result speaks in product terms ("Download the
    XLSX report") and never leaks engine class names.
"""
from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from django.conf import settings

from ai.engine.agent.plugins import ToolPlugin

logger = logging.getLogger("carbon.ai.plugins.export_document")

_SAFE_FMT = {"docx", "xlsx"}


def _slugify(value: str) -> str:
    value = re.sub(r"[^\w\s-]", "", value or "").strip().lower()
    value = re.sub(r"[-\s]+", "-", value)
    return value[:60] or "document"


class ExportDocument(ToolPlugin):
    name = "export_document"
    description = (
        "Generate a downloadable Word (.docx) and/or Excel (.xlsx) document "
        "from the conversation's findings — e.g. 'export this study as a Word "
        "report and an Excel comparison table'. Provide a title, optional "
        "markdown content, and/or a table (headers + rows). Returns a download "
        "link surfaced in chat."
    )
    input_schema: dict[str, Any] = {
        "type": "object",
        "properties": {
            "title": {
                "type": "string",
                "description": "Document title (used as heading and filename stem).",
            },
            "format": {
                "type": "string",
                "enum": ["docx", "xlsx", "both"],
                "description": "Which formats to generate. Default: both.",
            },
            "content": {
                "type": "string",
                "description": (
                    "Markdown body: '# Heading', '## Subheading', '- bullet', "
                    "'1. item', blank-line-separated paragraphs. Rendered into "
                    "the Word document (and a summary sheet for Excel)."
                ),
            },
            "table": {
                "type": "object",
                "description": "Optional structured table for Excel / Word.",
                "properties": {
                    "headers": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Column headers.",
                    },
                    "rows": {
                        "type": "array",
                        "items": {"type": "array", "items": {"type": "string"}},
                        "description": "Rows of cell values (strings/numbers).",
                    },
                },
                "required": ["headers", "rows"],
            },
        },
        "required": ["title"],
    }
    requires_confirmation = False
    capability: str | None = None
    app_identifier: str | None = None

    async def execute(self, args: dict, *, ctx) -> dict:
        title = (args.get("title") or "").strip()
        if not title:
            return {"error": "A title is required — e.g. 'Carbon Standards Study'."}

        fmt = (args.get("format") or "both").strip().lower()
        if fmt not in ("docx", "xlsx", "both"):
            fmt = "both"

        content = (args.get("content") or "").strip()
        table = args.get("table") or None

        out_dir = Path(settings.MEDIA_ROOT) / "ai_exports"
        try:
            out_dir.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            logger.exception("export_document mkdir failed")
            return {"error": f"Could not create export folder: {exc}"}

        stem = _slugify(title)
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        files: list[dict] = []

        if fmt in ("docx", "both"):
            filename = f"{stem}-{stamp}.docx"
            path = out_dir / filename
            try:
                self._write_docx(path, title, content, table)
                files.append({
                    "filename": filename,
                    "format": "docx",
                    "path": f"/media/ai_exports/{filename}",
                })
            except Exception as exc:  # fail-visible, never fabricate a file
                logger.exception("export_document docx failed")
                return {"error": f"Word export failed: {exc}"}

        if fmt in ("xlsx", "both"):
            filename = f"{stem}-{stamp}.xlsx"
            path = out_dir / filename
            try:
                self._write_xlsx(path, title, content, table)
                files.append({
                    "filename": filename,
                    "format": "xlsx",
                    "path": f"/media/ai_exports/{filename}",
                })
            except Exception as exc:
                logger.exception("export_document xlsx failed")
                return {"error": f"Excel export failed: {exc}"}

        if not files:
            return {"error": "No document was generated."}

        return {
            "requires_confirmation": False,
            "action": "download",
            "title": title,
            "files": files,
            "message": (
                f"Exported “{title}” as "
                + ", ".join(f.get("format", "").upper() for f in files)
                + ". Use the download button(s) below."
            ),
        }

    # ── generators ─────────────────────────────────────────────────────────

    def _write_docx(self, path: Path, title: str, content: str, table: dict | None) -> None:
        from docx import Document

        doc = Document()
        doc.add_heading(title, level=0)
        if content:
            self._render_markdown_to_docx(doc, content)
        if table:
            headers = table.get("headers") or []
            rows = table.get("rows") or []
            if headers:
                t = doc.add_table(rows=1, cols=len(headers))
                t.style = "Light Grid Accent 1"
                for i, h in enumerate(headers):
                    t.rows[0].cells[i].text = str(h)
                for row in rows:
                    cells = t.add_row().cells
                    for i in range(len(headers)):
                        cells[i].text = str(row[i]) if i < len(row) else ""
        doc.save(str(path))

    @staticmethod
    def _render_markdown_to_docx(doc, content: str) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif re.match(r"^[-*•]\s+", stripped):
                doc.add_paragraph(stripped[2:].lstrip(), style="List Bullet")
            elif re.match(r"^\d+[.)]\s+", stripped):
                doc.add_paragraph(re.sub(r"^\d+[.)]\s+", "", stripped), style="List Number")
            elif stripped.startswith("> "):
                p = doc.add_paragraph(stripped[2:])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = None
            else:
                doc.add_paragraph(stripped)

    def _write_xlsx(self, path: Path, title: str, content: str, table: dict | None) -> None:
        from openpyxl import Workbook
        from openpyxl.styles import Font

        wb = Workbook()
        ws = wb.active
        ws.title = "Report"
        ws["A1"] = title
        ws["A1"].font = Font(bold=True, size=14)

        row = 3
        if table:
            headers = table.get("headers") or []
            rows = table.get("rows") or []
            if headers:
                for c, h in enumerate(headers, start=1):
                    cell = ws.cell(row=row, column=c, value=str(h))
                    cell.font = Font(bold=True)
                row += 1
                for r in rows:
                    for c in range(len(headers)):
                        ws.cell(row=row, column=c + 1, value=r[c] if c < len(r) else "")
                    row += 1
                row += 1
        if content:
            ws.cell(row=row, column=1, value="Summary").font = Font(bold=True)
            row += 1
            for line in content.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                stripped = re.sub(r"^#+\s+", "", stripped)
                stripped = re.sub(r"^[-*•]\s+", "• ", stripped)
                ws.cell(row=row, column=1, value=stripped[:4000])
                row += 1

        wb.save(str(path))
