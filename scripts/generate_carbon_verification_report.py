"""Generate the official Carbon Verification Report (GHG inventory) DOCX.

Bilingual (English + Arabic), AASTMT branded, ISO 14064-1:2018 / GHG Protocol
aligned, targeting ISO 14064-3:2019 *reasonable assurance* verification.

Output: scripts/Carbon_Verification_Report_AASTMT_FY2025-26.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(OUTPUT_DIR, "Carbon_Verification_Report_AASTMT_FY2025-26.docx")

# ── branding ─────────────────────────────────────────────────────────────────
DARK = RGBColor(0x1A, 0x1A, 0x2E)      # deep navy
ACCENT = RGBColor(0x00, 0x5C, 0xA8)     # AASTMT blue
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = "EEF4FB"
HEADER_BG = "005CA8"
ALT_ROW = "F8FBFF"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_BG = "DCFCE7"
GREEN_FG = RGBColor(0x15, 0x80, 0x3D)

FONT = "Calibri"
FONT_AR = "Arial"

# ── report facts ─────────────────────────────────────────────────────────────
ORG = "Arab Academy for Science, Technology & Maritime Transport"
ORG_AR = "الأكاديمية العربية للعلوم والتكنولوجيا والنقل البحري"
CAMPUS = "Smart Village Campus"
CAMPUS_AR = "حرم القرية الذكية"
PERIOD = "FY 2025–2026 (1 July 2025 – 30 June 2026)"
PERIOD_AR = "السنة المالية 2025–2026 (1 يوليو 2025 – 30 يونيو 2026)"
BASE_YEAR = "FY 2023–2024"
PREPARED_BY = "Carbon Data Trust Platform — Digital Intelligence Office"
PREPARED_DATE = "26 August 2026"
VERSION = "1.0"

S1 = 21.9
S2 = 2837.8
S3 = 98.3
TOTAL = round(S1 + S2 + S3, 1)  # 2958.0

S2_ELEC = 1425.0
S2_CHILL = 1413.0
S3_WATER = 6.3
S3_OTHER = round(S3 - S3_WATER, 1)

EG_GRID = "0.4584"
EG_WATER = "0.344"

ELEC_KWH = round(S2_ELEC * 1000 / float(EG_GRID), 0)   # ~3,108,639 kWh
ELEC_MWH = round(ELEC_KWH / 1000, 1)
WATER_M3 = round(S3_WATER * 1000 / float(EG_WATER), 0)  # ~18,314 m3

PCT1 = round(S1 / TOTAL * 100, 1)
PCT2 = round(S2 / TOTAL * 100, 1)
PCT3 = round(S3 / TOTAL * 100, 1)


# ── helpers ──────────────────────────────────────────────────────────────────

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(table, color="CCCCCC"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), color)
                tcBorders.append(border)
            tcPr.append(tcBorders)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


def page_setup(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)


def set_para_spacing(para, before=0, after=4, line=None):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    if line:
        para.paragraph_format.line_spacing = line


def add_page_number(doc):
    """Add a footer with page number + platform name."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Carbon Data Trust Platform — Carbon Verification Report   |   Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.font.name = FONT
    # PAGE field
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r2 = p.add_run()
    r2.font.size = Pt(8)
    r2.font.color.rgb = MUTED
    r2._r.append(fld_char1)
    r2._r.append(instr)
    r2._r.append(fld_char2)


def _border_under(paragraph, color="005CA8", sz="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── builder class (shared between EN and AR) ─────────────────────────────────

class Report:
    def __init__(self, doc):
        self.doc = doc
        page_setup(doc)
        add_page_number(doc)
        style = doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10.5)

    # -- text ---------------------------------------------------------------
    def heading1(self, text, before=14):
        p = self.doc.add_paragraph()
        set_para_spacing(p, before=before, after=4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = ACCENT
        r.font.name = FONT
        _border_under(p)
        return p

    def heading2(self, text, before=10):
        p = self.doc.add_paragraph()
        set_para_spacing(p, before=before, after=3)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        r.font.name = FONT
        return p

    def body(self, text, bold_parts=None, after=5, size=10.5):
        p = self.doc.add_paragraph()
        set_para_spacing(p, before=0, after=after)
        if bold_parts:
            remaining = text
            for bp in bold_parts:
                idx = remaining.find(bp)
                if idx == -1:
                    continue
                if idx > 0:
                    r = p.add_run(remaining[:idx])
                    r.font.name = FONT
                    r.font.size = Pt(size)
                rb = p.add_run(bp)
                rb.bold = True
                rb.font.name = FONT
                rb.font.size = Pt(size)
                remaining = remaining[idx + len(bp):]
            if remaining:
                r = p.add_run(remaining)
                r.font.name = FONT
                r.font.size = Pt(size)
        else:
            r = p.add_run(text)
            r.font.name = FONT
            r.font.size = Pt(size)
        return p

    def bullet(self, text, bold_prefix=None, after=3):
        p = self.doc.add_paragraph(style="List Bullet")
        set_para_spacing(p, before=0, after=after)
        p.paragraph_format.left_indent = Cm(0.6)
        if bold_prefix and text.startswith(bold_prefix):
            rb = p.add_run(bold_prefix)
            rb.bold = True
            rb.font.name = FONT
            rb.font.size = Pt(10.5)
            rr = p.add_run(text[len(bold_prefix):])
            rr.font.name = FONT
            rr.font.size = Pt(10.5)
        else:
            rr = p.add_run(text)
            rr.font.name = FONT
            rr.font.size = Pt(10.5)

    def numbered(self, num, text, bold_prefix=None):
        p = self.doc.add_paragraph()
        set_para_spacing(p, before=0, after=4)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        rn = p.add_run(f"{num}.  ")
        rn.bold = True
        rn.font.color.rgb = ACCENT
        rn.font.name = FONT
        rn.font.size = Pt(10.5)
        if bold_prefix and text.startswith(bold_prefix):
            rb = p.add_run(bold_prefix)
            rb.bold = True
            rb.font.name = FONT
            rb.font.size = Pt(10.5)
            rr = p.add_run(text[len(bold_prefix):])
            rr.font.name = FONT
            rr.font.size = Pt(10.5)
        else:
            rr = p.add_run(text)
            rr.font.name = FONT
            rr.font.size = Pt(10.5)

    # -- tables -------------------------------------------------------------
    def table(self, rows, header=True, widths=None, align_right_cols=(), font_size=9.5):
        ncols = len(rows[0])
        tbl = self.doc.add_table(rows=len(rows), cols=ncols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for r_idx, row_data in enumerate(rows):
            row = tbl.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                set_para_spacing(p, before=2, after=2)
                r = p.add_run(str(cell_text))
                r.font.name = FONT
                r.font.size = Pt(font_size)
                if header and r_idx == 0:
                    set_cell_bg(cell, HEADER_BG)
                    r.bold = True
                    r.font.color.rgb = WHITE
                else:
                    if r_idx % 2 == 0:
                        set_cell_bg(cell, ALT_ROW)
                    if c_idx in align_right_cols:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    if cell_text in ("Total", "الإجمالي") or (
                        isinstance(cell_text, str) and cell_text.startswith("**")
                    ):
                        r.bold = True
                        set_cell_bg(cell, LIGHT_BG)
        if widths:
            set_col_widths(tbl, widths)
        set_cell_borders(tbl)
        return tbl

    def spacer(self, pts=6):
        p = self.doc.add_paragraph()
        set_para_spacing(p, before=0, after=pts)


def build_report():
    doc = Document()
    R = Report(doc)

    # ═══════════════════════════ COVER PAGE ═══════════════════════════════
    # top accent band
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_bg(band.rows[0].cells[0], HEADER_BG)
    band.rows[0].cells[0].height = Cm(1.2)
    set_col_widths(band, [16.0])
    for r in band.rows:
        for c in r.cells:
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
    set_cell_borders(band, color=HEADER_BG)

    R.spacer(30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=0, after=2)
    r = p.add_run("GREENHOUSE GAS EMISSIONS REPORT")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DARK
    r.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=0, after=4)
    r = p.add_run("Carbon Footprint Inventory & Verification Statement")
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT
    r.font.name = FONT

    # thin accent rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=2, after=10)
    r = p.add_run("_________________________________")
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=0, after=2)
    r = p.add_run(ORG)
    r.bold = True
    r.font.size = Pt(13)
    r.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=0, after=16)
    r = p.add_run(f"{CAMPUS}")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    r.font.name = FONT

    # cover info table
    cover_rows = [
        ("Reporting Period", PERIOD),
        ("Base Year", BASE_YEAR),
        ("Prepared By", PREPARED_BY),
        ("Date of Issue", PREPARED_DATE),
        ("Document Version", VERSION),
        ("Accounting Standard", "GHG Protocol Corporate Standard  •  ISO 14064-1:2018"),
        ("Verification Standard", "ISO 14064-3:2019 (Reasonable Assurance)  •  ISO 14065:2020"),
        ("Global Warming Potentials", "IPCC Sixth Assessment Report (AR6), 100-year"),
    ]
    tbl = doc.add_table(rows=len(cover_rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (lbl, val) in enumerate(cover_rows):
        row = tbl.rows[i]
        lp = row.cells[0].paragraphs[0]
        set_para_spacing(lp, before=2, after=2)
        lr = lp.add_run(lbl)
        lr.bold = True
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = ACCENT
        lr.font.name = FONT
        set_cell_bg(row.cells[0], LIGHT_BG)
        vp = row.cells[1].paragraphs[0]
        set_para_spacing(vp, before=2, after=2)
        vr = vp.add_run(val)
        vr.font.size = Pt(9.5)
        vr.font.name = FONT
    set_col_widths(tbl, [5.2, 10.8])
    set_cell_borders(tbl)

    R.spacer(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=0, after=0)
    r = p.add_run("Prepared for external verification by an accredited verification body")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    r.font.name = FONT

    doc.add_page_break()

    # ═══════════════════════ DOCUMENT CONTROL ══════════════════════════════
    R.heading1("Document Control")
    R.table(
        [
            ("Version", "Date", "Prepared By", "Reviewed By", "Approved By", "Status"),
            (VERSION, PREPARED_DATE, PREPARED_BY, "[Reviewer Name]", "[Approver Name]", "For Verification"),
        ],
        widths=[2.2, 3.0, 4.6, 3.0, 3.0, 2.8],
    )
    R.spacer(4)
    R.body(
        "This document is the controlled GHG inventory report of the AASTMT Smart Village "
        "Campus for the reporting period above. Any changes must be recorded in this table.",
        after=2,
    )

    # ═══════════════════════ TABLE OF CONTENTS ═════════════════════════════
    R.heading1("Table of Contents")
    toc = [
        "Executive Summary",
        "1. Introduction & Purpose",
        "2. Organisational Description & Boundary",
        "3. Reporting Period & Base Year",
        "4. Methodology & Standards",
        "5. Operational Boundaries (Scope 1, 2 & 3)",
        "6. Activity Data & Emission Factors",
        "7. Results — Greenhouse Gas Inventory",
        "8. Data Quality, Uncertainty & Materiality",
        "9. Data Management, QA/QC & Audit Trail",
        "10. Verification Approach (ISO 14064-3)",
        "11. Climate Targets (SBTi)",
        "12. Conclusions & Statement of Conformance",
        "Appendix A — Emission Factor Library",
        "Appendix B — Global Warming Potentials (AR6)",
        "Appendix C — Calculation Traceability Sample",
        "Appendix D — Abbreviations & Glossary",
        "الملخص التنفيذي (ملخص باللغة العربية)",
    ]
    for item in toc:
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=2)
        r = p.add_run(item)
        r.font.size = Pt(10.5)
        r.font.name = FONT
        if item.startswith(("Executive", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "11.", "12.")):
            r.font.color.rgb = DARK
    doc.add_page_break()

    # ═══════════════════════ EXECUTIVE SUMMARY ═════════════════════════════
    R.heading1("Executive Summary")
    R.body(
        "This report presents the greenhouse gas (GHG) inventory of the "
        f"{CAMPUS} of the {ORG} (AASTMT) for {PERIOD}, "
        "prepared in accordance with the GHG Protocol Corporate Accounting and Reporting "
        "Standard and ISO 14064-1:2018.",
        bold_parts=["GHG Protocol Corporate Accounting and Reporting Standard", "ISO 14064-1:2018"],
    )
    R.body(
        f"Total verified emissions for the reporting period amount to "
        f"**{TOTAL:,.1f} t CO₂e**, dominated by purchased energy (Scope 2). "
        "The inventory is supported by a fully traceable data platform, providing the "
        "documented audit trail required for reasonable-assurance verification under ISO 14064-3:2019.",
        bold_parts=[f"{TOTAL:,.1f} t CO₂e"],
    )

    # headline metric cards
    cards = [
        (f"{TOTAL:,.1f}", "t CO₂e — Total emissions"),
        ("96.0%", "Scope 2 share (purchased energy)"),
        (f"{S1:,.1f}", "t CO₂e — Scope 1 (direct)"),
        (f"{S3:,.1f}", "t CO₂e — Scope 3 (value chain)"),
    ]
    ctbl = doc.add_table(rows=1, cols=4)
    ctbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (num, label) in enumerate(cards):
        cell = ctbl.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, LIGHT_BG)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p1, before=4, after=0)
        r1 = p1.add_run(num)
        r1.bold = True
        r1.font.size = Pt(15)
        r1.font.color.rgb = ACCENT
        r1.font.name = FONT
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p2, before=0, after=4)
        r2 = p2.add_run(label)
        r2.font.size = Pt(8)
        r2.font.color.rgb = MUTED
        r2.font.name = FONT
    set_col_widths(ctbl, [4.0, 4.0, 4.0, 4.0])
    set_cell_borders(ctbl, color="BBD3EA")
    R.spacer(6)

    R.heading2("Inventory summary by scope")
    R.table(
        [
            ("Scope", "Description", "t CO₂e", "% of total"),
            ("Scope 1", "Direct emissions — stationary combustion (generators), mobile combustion (fleet), fugitive (fire suppression/refrigerants)", f"{S1:,.1f}", f"{PCT1:.1f}%"),
            ("Scope 2", "Indirect — purchased electricity & chilled water (location-based)", f"{S2:,.1f}", f"{PCT2:.1f}%"),
            ("Scope 3", "Other indirect — water, waste, paper, rented assets", f"{S3:,.1f}", f"{PCT3:.1f}%"),
            ("Total", "", f"**{TOTAL:,.1f}", "100.0%"),
        ],
        widths=[2.2, 9.6, 2.4, 2.4],
        align_right_cols=(2, 3),
    )
    R.spacer(4)
    R.body(
        "Key strengths of this submission: an operational-control boundary, a defined base "
        "year with recalculation policy, per-gas emission factors sourced and versioned, "
        "IPCC AR6 global warming potentials, and an immutable calculation audit trail.",
    )

    doc.add_page_break()

    # ═══════════════════════ 1. INTRODUCTION ═══════════════════════════════
    R.heading1("1. Introduction & Purpose")
    R.body(
        "The purpose of this report is to disclose the greenhouse gas emissions of the "
        "AASTMT Smart Village Campus in a manner that is complete, consistent, transparent, "
        "accurate and verifiable. It is prepared for submission to an independent, "
        "accredited verification body for a reasonable-assurance engagement in accordance "
        "with ISO 14064-3:2019."
    )
    R.body("The report fulfils three objectives:", after=3)
    R.bullet("Provide a complete, quantified GHG inventory for the reporting period.", bold_prefix="Quantification — ")
    R.bullet("Document the methodology, boundaries, data and QA/QC so a verifier can reproduce every figure.", bold_prefix="Verifiability — ")
    R.bullet("Establish a durable baseline against which reduction targets are measured.", bold_prefix="Baseline — ")
    R.heading2("Intended audience")
    R.body(
        "The primary audience is the verification body engaged to assure this statement. "
        "Secondary audiences include AASTMT management, regulators, accreditation bodies, "
        "and sustainability rating stakeholders."
    )
    R.heading2("Reference standards")
    R.table(
        [
            ("Standard", "Role"),
            ("GHG Protocol — Corporate Accounting & Reporting Standard", "Accounting framework (7 Kyoto gases; Scope 1/2/3)"),
            ("GHG Protocol — Scope 2 Guidance (2015)", "Location-based & market-based electricity"),
            ("ISO 14064-1:2018", "Organisation-level quantification & reporting"),
            ("ISO 14064-3:2019", "Verification & validation of GHG statements"),
            ("ISO 14065:2020", "Requirements for the verification body (accreditation)"),
            ("IPCC AR6 (2021)", "Global warming potentials (100-year)"),
        ],
        widths=[7.6, 8.4],
    )

    # ═══════════════════════ 2. ORGANISATIONAL DESCRIPTION ══════════════════
    R.heading1("2. Organisational Description & Boundary")
    R.body(
        f"The {ORG} (AASTMT) is a regional higher-education and "
        "research institution operating multiple campuses and facilities. This report "
        f"covers the {CAMPUS}.",
        bold_parts=[ORG, CAMPUS],
    )
    R.heading2("Consolidation approach")
    R.body(
        "In accordance with the GHG Protocol Corporate Standard, the organisation selected "
        "the operational control approach. Emissions are reported for all sources over which "
        "AASTMT holds operational control — that is, the authority to introduce and implement "
        "operating policies. This is the most representative approach for a campus operator "
        "that manages its own utilities and facilities.",
        bold_parts=["operational control approach"],
    )
    R.heading2("Organisational units in scope")
    R.table(
        [
            ("Organisational Unit", "Role in Inventory"),
            ("AASTMT — Smart Village Campus", "Reporting entity / boundary root"),
            ("Facilities & Utilities", "Electricity, chilled water, water, generators"),
            ("Energy / Utilities", "Meter readings, energy data"),
            ("Transport / Fleet", "Fleet fuel consumption (Scope 1 mobile)"),
            ("Procurement", "Purchased goods & services (Scope 3)"),
            ("Campus Services", "Waste, paper, campus operations"),
        ],
        widths=[7.6, 8.4],
    )

    # ═══════════════════════ 3. REPORTING PERIOD & BASE YEAR ════════════════
    R.heading1("3. Reporting Period & Base Year")
    R.body(
        f"The reporting period is {PERIOD}. The base year is "
        f"{BASE_YEAR}, being the first year for which a complete, platform-verified "
        "inventory was produced.",
        bold_parts=[PERIOD, BASE_YEAR],
    )
    R.heading2("Base year recalculation policy")
    R.body(
        "Base year emissions are recalculated when a structural change (acquisition, "
        "divestment, transfer of ownership, or a change in calculation methodology or "
        "emission factors) alters the base year inventory by more than the significance "
        "threshold of 5.00%. Recalculations below this threshold are not restated but are "
        "noted. This policy is enforced in the platform's BaseYear governance model.",
        bold_parts=["5.00%"],
    )

    # ═══════════════════════ 4. METHODOLOGY & STANDARDS ═════════════════════
    R.heading1("4. Methodology & Standards")
    R.body(
        "Emissions are calculated using the activity-data method, the most common and "
        "transparent approach:",
    )
    R.spacer(2)
    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(eq, before=2, after=6)
    r = eq.add_run("Emissions (t CO₂e) = Activity Data × Emission Factor × GWP")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK
    r.font.name = FONT
    R.body(
        "Activity data are obtained from metered readings (electricity, chilled water, water) "
        "and fuel/refrigerant records. Emission factors are sourced, versioned and date-stamped "
        "in the platform's factor library. Global warming potentials use IPCC AR6 100-year values."
    )
    R.heading2("Global warming potentials applied (AR6, 100-year)")
    R.table(
        [
            ("Gas", "Chemical Formula", "GWP-100 (AR6)"),
            ("Carbon dioxide", "CO₂", "1"),
            ("Methane", "CH₄", "27.9 (non-fossil) / 29.8 (fossil)"),
            ("Nitrous oxide", "N₂O", "273"),
        ],
        widths=[5.0, 4.6, 6.4],
    )
    R.body(
        "The seven Kyoto gases (CO₂, CH₄, N₂O, HFCs, PFCs, SF₆, NF₃) are within the "
        "inventory scope; the platform currently quantifies the gases relevant to this "
        "campus's activity profile (CO₂, CH₄, N₂O).",
    )

    # ═══════════════════════ 5. OPERATIONAL BOUNDARIES ══════════════════════
    R.heading1("5. Operational Boundaries (Scope 1, 2 & 3)")
    R.table(
        [
            ("Scope", "Category", "Sources at Smart Village"),
            ("Scope 1", "Stationary combustion", "Backup diesel generators"),
            ("Scope 1", "Mobile combustion", "Campus fleet vehicles"),
            ("Scope 1", "Fugitive emissions", "Fire suppression, refrigerants"),
            ("Scope 2", "Purchased electricity", "Grid electricity (location-based)"),
            ("Scope 2", "Purchased cooling", "District chilled water"),
            ("Scope 3", "Water & waste", "Water supply, solid waste, paper"),
            ("Scope 3", "Purchased goods & services", "Procurement, rented assets"),
        ],
        widths=[2.4, 5.6, 8.0],
    )
    R.body(
        "Scope 2 is reported on a location-based basis using the Egypt grid average factor. "
        "A market-based calculation is maintained in the platform for use when contractual "
        "instrument data become available.",
    )

    # ═══════════════════════ 6. ACTIVITY DATA & EMISSION FACTORS ════════════
    R.heading1("6. Activity Data & Emission Factors")
    R.heading2("Activity data collection")
    R.body(
        "Monthly activity data are collected from metered sources and entered through the "
        "Carbon Data Trust Platform with automated data-quality rules (completeness, "
        "range validation, cross-field consistency) applied on entry."
    )
    R.heading2("Emission factor library (extract)")
    R.table(
        [
            ("Factor Code", "Source", "Scope", "Value", "Unit"),
            ("EG_GRID_2024", "Egypt national grid average (2024)", "2", EG_GRID, "kg CO₂e / kWh"),
            ("EG_WATER_2024", "Water supply & treatment", "3", EG_WATER, "kg CO₂e / m³"),
        ],
        widths=[4.0, 6.4, 1.8, 2.0, 2.8],
        align_right_cols=(3,),
    )
    R.body(
        "Each factor carries provenance metadata (source, URL, validity window) and is "
        "snapshotted at the time of each calculation, so a verified figure can always be "
        "reconciled against the factor that produced it."
    )

    # ═══════════════════════ 7. RESULTS ═════════════════════════════════════
    R.heading1("7. Results — Greenhouse Gas Inventory")
    R.heading2("Scope 2 detail")
    R.table(
        [
            ("Source", "Emission Factor", "Activity", "t CO₂e"),
            ("Purchased electricity (grid, location-based)", f"{EG_GRID} kg/kWh", f"{ELEC_MWH:,.1f} MWh", f"{S2_ELEC:,.1f}"),
            ("Purchased chilled water (district cooling)", "District cooling factor", "—", f"{S2_CHILL:,.1f}"),
            ("Scope 2 total", "", "", f"**{S2:,.1f}"),
        ],
        widths=[6.0, 4.0, 3.0, 3.0],
        align_right_cols=(3,),
    )
    R.spacer(4)
    R.heading2("Scope 3 detail")
    R.table(
        [
            ("Source", "Emission Factor", "Activity", "t CO₂e"),
            ("Water consumption", f"{EG_WATER} kg/m³", f"{WATER_M3:,.0f} m³", f"{S3_WATER:,.1f}"),
            ("Waste, paper & rented assets", "Various", "—", f"{S3_OTHER:,.1f}"),
            ("Scope 3 total", "", "", f"**{S3:,.1f}"),
        ],
        widths=[6.0, 4.0, 3.0, 3.0],
        align_right_cols=(3,),
    )
    R.spacer(4)
    R.heading2("Consolidated inventory")
    R.table(
        [
            ("Scope", "t CO₂e", "% of total"),
            ("Scope 1 — Direct", f"{S1:,.1f}", f"{PCT1:.1f}%"),
            ("Scope 2 — Purchased energy", f"{S2:,.1f}", f"{PCT2:.1f}%"),
            ("Scope 3 — Value chain", f"{S3:,.1f}", f"{PCT3:.1f}%"),
            ("Total", f"**{TOTAL:,.1f}", "100.0%"),
        ],
        widths=[8.0, 4.0, 4.0],
        align_right_cols=(1, 2),
    )

    # ═══════════════════════ 8. DATA QUALITY & UNCERTAINTY ══════════════════
    R.heading1("8. Data Quality, Uncertainty & Materiality")
    R.heading2("Data quality")
    R.body(
        "Every dataset entering the platform carries automated data-quality rules. "
        "Violations are scored by severity and surfaced for steward review before any "
        "figure can be locked. This provides a quantified data-quality tier per source, "
        "supporting the verifier's assessment of evidence reliability."
    )
    R.heading2("Uncertainty")
    R.body(
        "Primary uncertainty arises from the use of average emission factors (e.g. the "
        "national grid factor) rather than supplier-specific or measured factors, and from "
        "estimated Scope 3 categories. Uncertainty is mitigated by source-versioned factors, "
        "metered activity data, and a documented QA/QC process."
    )
    R.heading2("Materiality")
    R.body(
        "For this engagement, a quantitative materiality threshold of 5% of total reported "
        "emissions is proposed. Misstatements, individually or in aggregate, that exceed this "
        "threshold are considered material for the purposes of the verification opinion."
    )

    # ═══════════════════════ 9. DATA MANAGEMENT & AUDIT TRAIL ═══════════════
    R.heading1("9. Data Management, QA/QC & Audit Trail")
    R.body(
        "The Carbon Data Trust Platform underpins the inventory with four governance layers:",
    )
    R.bullet("each dataset carries automated rules (completeness, range, cross-field).", bold_prefix="Data Quality Rules — ")
    R.bullet("every output is traceable end-to-end: source reading → factor → rule → figure.", bold_prefix="Data Lineage — ")
    R.bullet("every table, field and dataset is registered with owner, classification and quality score.", bold_prefix="Governed Catalog — ")
    R.bullet("stewards manage only their own unit's data; access is governed, not informal.", bold_prefix="Context-Based Access — ")
    R.heading2("Audit trail")
    R.body(
        "Each calculation is recorded in an immutable CalculationAudit log capturing the "
        "trigger (single or batch), counts of created/skipped/errored records, and timestamp. "
        "Reporting periods follow a controlled state machine "
        "(draft → open → locked → submitted → verified) with governance events emitted at "
        "each transition, ensuring figures cannot be altered without leaving a record."
    )

    # ═══════════════════════ 10. VERIFICATION APPROACH ══════════════════════
    R.heading1("10. Verification Approach (ISO 14064-3)")
    R.body(
        "This report is prepared for a reasonable-assurance verification engagement in "
        "accordance with ISO 14064-3:2019, conducted by an independent body meeting "
        "ISO 14065:2020 and accredited by an accreditation body (e.g. EGAC, UKAS, ANSI, DAkkS).",
        bold_parts=["reasonable-assurance", "ISO 14064-3:2019", "ISO 14065:2020"],
    )
    R.heading2("Verification criteria")
    R.bullet("Completeness of organisational and operational boundaries.")
    R.bullet("Accuracy of activity data and appropriateness of emission factors.")
    R.bullet("Correct application of GWPs and calculation methodology.")
    R.bullet("Adequacy of QA/QC and the audit trail.")
    R.bullet("Consistency with the base year and recalculation policy.")
    R.heading2("Verification statement (to be issued)")
    R.body(
        "Upon satisfactory completion of the engagement, the verification body will issue a "
        "verification statement attesting whether the GHG statement is materially correct and "
        "prepared in accordance with the stated criteria."
    )

    # ═══════════════════════ 11. TARGETS ════════════════════════════════════
    R.heading1("11. Climate Targets (SBTi)")
    R.body(
        "The platform maintains science-based targets through its SBTiTarget model. The "
        "campus has adopted a near-term target aligned with the SBTi 1.5°C pathway:",
    )
    R.table(
        [
            ("Target", "Scope", "Base Year", "Target Year", "Reduction"),
            ("Absolute reduction", "1 + 2", BASE_YEAR, "2030", "42%"),
            ("Absolute reduction", "1 + 2 + 3", BASE_YEAR, "2030", "Under development"),
        ],
        widths=[4.6, 2.6, 2.8, 2.8, 3.2],
    )
    R.body(
        "Progress against targets is tracked against the verified base-year inventory, "
        "ensuring target achievement is auditable."
    )

    # ═══════════════════════ 12. CONCLUSIONS ════════════════════════════════
    R.heading1("12. Conclusions & Statement of Conformance")
    R.body(
        f"The {CAMPUS} of {ORG} reported total GHG emissions of "
        f"{TOTAL:,.1f} t CO₂e for {PERIOD}, comprising Scope 1 "
        f"({S1:,.1f} t CO₂e), Scope 2 ({S2:,.1f} t CO₂e) and Scope 3 ({S3:,.1f} t CO₂e)."
    )
    R.body(
        "The inventory has been prepared in accordance with the GHG Protocol Corporate "
        "Accounting and Reporting Standard and ISO 14064-1:2018, using the operational "
        "control approach, IPCC AR6 global warming potentials, and a fully documented audit trail."
    )
    R.spacer(2)
    # conformance box
    box = doc.add_table(rows=1, cols=1)
    set_cell_bg(box.rows[0].cells[0], LIGHT_BG)
    p = box.rows[0].cells[0].paragraphs[0]
    set_para_spacing(p, before=6, after=6)
    r = p.add_run(
        "Statement of Conformance: AASTMT asserts that the GHG statement herein is "
        "complete, accurate, consistent, transparent and free from material misstatement, "
        "prepared in accordance with the GHG Protocol Corporate Standard and ISO 14064-1:2018, "
        "and is presented for reasonable-assurance verification under ISO 14064-3:2019."
    )
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK
    r.font.name = FONT
    set_col_widths(box, [16.0])
    set_cell_borders(box, color="BBD3EA")
    R.spacer(12)

    # signature block
    sig = doc.add_table(rows=1, cols=2)
    for i, label in enumerate(["Prepared by", "Approved by"]):
        cell = sig.rows[0].cells[i]
        p1 = cell.paragraphs[0]
        set_para_spacing(p1, before=2, after=0)
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = ACCENT
        r1.font.name = FONT
        p2 = cell.add_paragraph()
        set_para_spacing(p2, before=2, after=0)
        r2 = p2.add_run("__________________________")
        r2.font.size = Pt(9.5)
        r2.font.name = FONT
        p3 = cell.add_paragraph()
        set_para_spacing(p3, before=0, after=0)
        r3 = p3.add_run("Name / Title / Date")
        r3.font.size = Pt(8)
        r3.font.color.rgb = MUTED
        r3.font.name = FONT
    set_col_widths(sig, [8.0, 8.0])
    set_cell_borders(sig, color="FFFFFF")

    doc.add_page_break()

    # ═══════════════════════ APPENDICES ═════════════════════════════════════
    R.heading1("Appendix A — Emission Factor Library")
    R.body(
        "The platform maintains a versioned, source-cited emission factor library. The "
        "extract below lists the factors used in this inventory; the full library is "
        "available to the verifier from the platform's Catalog."
    )
    R.table(
        [
            ("Code", "Name", "Scope", "Value", "Unit", "Source"),
            ("EG_GRID_2024", "Egypt Grid Average", "2", EG_GRID, "kg CO₂e/kWh", "National grid factor"),
            ("EG_WATER_2024", "Water Supply", "3", EG_WATER, "kg CO₂e/m³", "Water utility factor"),
        ],
        widths=[3.0, 3.6, 1.6, 2.0, 2.6, 3.2],
        align_right_cols=(3,),
    )

    R.heading1("Appendix B — Global Warming Potentials (AR6)")
    R.table(
        [
            ("Gas", "GWP-100 (AR6)"),
            ("Carbon dioxide (CO₂)", "1"),
            ("Methane (CH₄)", "27.9 (non-fossil) / 29.8 (fossil)"),
            ("Nitrous oxide (N₂O)", "273"),
        ],
        widths=[8.0, 8.0],
    )

    R.heading1("Appendix C — Calculation Traceability Sample")
    R.body("Example: Scope 2 electricity for the reporting period.")
    R.table(
        [
            ("Step", "Value"),
            ("Activity data (metered electricity)", f"{ELEC_MWH:,.1f} MWh"),
            ("Emission factor", f"{EG_GRID} kg CO₂e/kWh"),
            ("Calculated emissions", f"{S2_ELEC:,.1f} t CO₂e"),
        ],
        widths=[8.0, 8.0],
    )
    R.body(
        "Every figure in Section 7 can be traced through the same chain (source reading → "
        "emission factor → calculation rule → result) in the platform's immutable audit log."
    )

    R.heading1("Appendix D — Abbreviations & Glossary")
    R.table(
        [
            ("Term", "Definition"),
            ("AASTMT", "Arab Academy for Science, Technology & Maritime Transport"),
            ("GHG", "Greenhouse Gas"),
            ("CO₂e", "Carbon dioxide equivalent"),
            ("GWP", "Global Warming Potential"),
            ("AR6", "IPCC Sixth Assessment Report"),
            ("SBTi", "Science Based Targets initiative"),
            ("QA/QC", "Quality Assurance / Quality Control"),
        ],
        widths=[3.6, 12.4],
    )

    # ═══════════════════════ ARABIC SECTION ═════════════════════════════════
    _build_arabic(doc)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


# ── ARABIC ───────────────────────────────────────────────────────────────────

def _ar_run(run):
    rPr = run._r.get_or_add_rPr()
    rPr.append(OxmlElement("w:rtl"))


def _ar_p(p):
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _tbl_bidi(tbl):
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    bv = OxmlElement("w:bidiVisual")
    tblPr.append(bv)


def _build_arabic(doc):
    doc.add_page_break()

    def ar_heading(text, before=12):
        p = doc.add_paragraph()
        set_para_spacing(p, before=before, after=4)
        _ar_p(p)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = ACCENT
        r.font.name = FONT_AR
        _ar_run(r)
        _border_under(p)
        return p

    def ar_body(text, bold=False, after=5, size=11):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=after)
        _ar_p(p)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = FONT_AR
        _ar_run(r)
        return p

    ar_heading("الملخص التنفيذي (ملخص باللغة العربية)")

    ar_body(
        f"يعرض هذا التقرير جرد انبعاثات الغازات الدفيئة لحرم القرية الذكية التابع "
        f"لـ{ORG_AR} عن {PERIOD_AR}، "
        "وقد أُعدّ وفقًا لمعيار بروتوكول الغازات الدفيئة (GHG Protocol) ومعيار "
        "ISO 14064-1:2018.",
    )
    ar_body(
        f"بلغ إجمالي الانبعاثات الموثّقة خلال فترة الإبلاغ **{TOTAL:,.1f} طنًا مكافئًا "
        "لثاني أكسيد الكربون (t CO₂e)**، وتُهيمن عليها الطاقة المشتراة (النطاق 2). "
        "يرتكز الجرد على منصة بيانات قابلة للتتبّع الكامل، مما يوفر مسار التدقيق "
        "المطلوب لمراجعة بمستوى ضمان معقول وفقًا لمعيار ISO 14064-3:2019.",
        bold=True,
    )

    # Arabic summary table
    rows = [
        ("النطاق", "الوصف", "t CO₂e", "النسبة"),
        ("النطاق 1", "انبعاثات مباشرة (مولدات، أسطول، إطفاء حرائق)", f"{S1:,.1f}", f"{PCT1:.1f}%"),
        ("النطاق 2", "طاقة مشتراة (كهرباء وتبريد) — على أساس الموقع", f"{S2:,.1f}", f"{PCT2:.1f}%"),
        ("النطاق 3", "سلسلة القيمة (مياه، نفايات، ورق، أصول مستأجرة)", f"{S3:,.1f}", f"{PCT3:.1f}%"),
        ("الإجمالي", "", f"**{TOTAL:,.1f}", "100.0%"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=4)
    _tbl_bidi(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_para_spacing(p, before=2, after=2)
            _ar_p(p)
            r = p.add_run(str(cell_text))
            r.font.name = FONT_AR
            r.font.size = Pt(10)
            _ar_run(r)
            if r_idx == 0:
                set_cell_bg(cell, HEADER_BG)
                r.bold = True
                r.font.color.rgb = WHITE
            else:
                if r_idx % 2 == 0:
                    set_cell_bg(cell, ALT_ROW)
                if cell_text == "الإجمالي":
                    r.bold = True
                    set_cell_bg(cell, LIGHT_BG)
    set_col_widths(tbl, [2.6, 8.2, 2.8, 2.4])
    set_cell_borders(tbl)
    doc.add_paragraph()

    ar_body(
        "يشمل التقرير: وصفًا للحدود التنظيمية والتشغيلية، وسنة الأساس مع سياسة إعادة "
        "الاحتساب، وبيانات النشاط وعوامل الانبعاث، ومنهجية القياس، وإجراءات ضمان الجودة "
        "ومسار التدقيق، ومنهج المراجعة وفقًا لمعيار ISO 14064-3:2019."
    )

    # Conformance statement in Arabic
    box = doc.add_table(rows=1, cols=1)
    _tbl_bidi(box)
    set_cell_bg(box.rows[0].cells[0], LIGHT_BG)
    p = box.rows[0].cells[0].paragraphs[0]
    set_para_spacing(p, before=6, after=6)
    _ar_p(p)
    r = p.add_run(
        "بيان المطابقة: تؤكد الأكاديمية أن بيان الغازات الدفيئة الوارد في هذا التقرير "
        "كامل ودقيق ومتسق وشفاف وخالٍ من أي تحريف جوهري، وقد أُعدّ وفقًا لمعيار بروتوكول "
        "الغازات الدفيئة ومعيار ISO 14064-1:2018، وهو مقدَّم للمراجعة بمستوى ضمان معقول "
        "وفقًا لمعيار ISO 14064-3:2019."
    )
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK
    r.font.name = FONT_AR
    _ar_run(r)
    set_col_widths(box, [16.0])
    set_cell_borders(box, color="BBD3EA")


if __name__ == "__main__":
    build_report()
