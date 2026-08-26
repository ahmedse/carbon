"""Generate chairman memo in English and Arabic DOCX formats."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ── helpers ──────────────────────────────────────────────────────────────────

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "CCCCCC")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def page_setup(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)


def set_para_spacing(para, before=0, after=4):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


DARK = RGBColor(0x1A, 0x1A, 0x2E)   # deep navy
ACCENT = RGBColor(0x00, 0x5C, 0xA8)  # AASTMT blue
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = "EEF4FB"
HEADER_BG = "005CA8"


# ── ENGLISH ──────────────────────────────────────────────────────────────────

def build_english():
    doc = Document()
    page_setup(doc)

    # ── default style ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)

    # ── top rule ─────────────────────────────────────────────────────────────
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(6)
    r = rule.add_run()
    r.font.color.rgb = ACCENT
    r.add_break()  # blank placeholder; the border below acts as the rule

    # ── title block ──────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    set_para_spacing(t, before=0, after=2)
    run = t.add_run("MEMORANDUM")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT
    run.font.name = "Calibri"

    sub = doc.add_paragraph()
    set_para_spacing(sub, before=0, after=10)
    sr = sub.add_run("Arab Academy for Science, Technology & Maritime Transport")
    sr.font.size = Pt(9)
    sr.font.color.rgb = MUTED
    sr.font.name = "Calibri"

    # ── header table ─────────────────────────────────────────────────────────
    htbl = doc.add_table(rows=4, cols=2)
    htbl.style = "Table Grid"
    htbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    labels = ["To", "From", "Date", "Subject"]
    values = [
        "H.E. Prof. [Chairman's Name], Chairman, Board of Trustees",
        "[Your Name], Digital Intelligence Office",
        "26 August 2026",
        "Institutional Data Platform — What We Built and What Comes Next",
    ]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        row = htbl.rows[i]
        row.cells[0].width = Cm(2.5)
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(lbl)
        lr.bold = True
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = ACCENT
        lr.font.name = "Calibri"
        set_cell_bg(row.cells[0], LIGHT_BG)

        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(val)
        vr.font.size = Pt(9.5)
        vr.font.name = "Calibri"
        if lbl == "Subject":
            vr.bold = True

    set_cell_borders(htbl)
    doc.add_paragraph()  # spacer

    # ── section helper ───────────────────────────────────────────────────────
    def section_heading(text):
        p = doc.add_paragraph()
        set_para_spacing(p, before=10, after=3)
        r2 = p.add_run(text.upper())
        r2.bold = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = ACCENT
        r2.font.name = "Calibri"
        # bottom border on heading
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "005CA8")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def body_para(text, bold_parts=None):
        """Add a body paragraph. bold_parts = list of substrings to bold."""
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=5)
        if bold_parts:
            remaining = text
            for bp in bold_parts:
                idx = remaining.find(bp)
                if idx == -1:
                    continue
                if idx > 0:
                    r3 = p.add_run(remaining[:idx])
                    r3.font.name = "Calibri"
                    r3.font.size = Pt(10.5)
                r3b = p.add_run(bp)
                r3b.bold = True
                r3b.font.name = "Calibri"
                r3b.font.size = Pt(10.5)
                remaining = remaining[idx + len(bp):]
            if remaining:
                r3c = p.add_run(remaining)
                r3c.font.name = "Calibri"
                r3c.font.size = Pt(10.5)
        else:
            r3 = p.add_run(text)
            r3.font.name = "Calibri"
            r3.font.size = Pt(10.5)
        return p

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        set_para_spacing(p, before=0, after=3)
        p.paragraph_format.left_indent = Cm(0.6)
        if bold_prefix and text.startswith(bold_prefix):
            rb = p.add_run(bold_prefix)
            rb.bold = True
            rb.font.name = "Calibri"
            rb.font.size = Pt(10.5)
            rr = p.add_run(text[len(bold_prefix):])
            rr.font.name = "Calibri"
            rr.font.size = Pt(10.5)
        else:
            rr = p.add_run(text)
            rr.font.name = "Calibri"
            rr.font.size = Pt(10.5)

    def numbered(num, text, bold_prefix=None):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=4)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        rn = p.add_run(f"{num}.  ")
        rn.bold = True
        rn.font.color.rgb = ACCENT
        rn.font.name = "Calibri"
        rn.font.size = Pt(10.5)
        if bold_prefix and text.startswith(bold_prefix):
            rb = p.add_run(bold_prefix)
            rb.bold = True
            rb.font.name = "Calibri"
            rb.font.size = Pt(10.5)
            rr = p.add_run(text[len(bold_prefix):])
            rr.font.name = "Calibri"
            rr.font.size = Pt(10.5)
        else:
            rr = p.add_run(text)
            rr.font.name = "Calibri"
            rr.font.size = Pt(10.5)

    # ── SECTION 1: The situation ──────────────────────────────────────────────
    section_heading("The Situation")
    body_para(
        "AASTMT runs several institutional systems — carbon reporting, faculty performance, "
        "facilities, research — each managed separately. Reporting across them takes weeks, "
        "data quality is inconsistent, and there is no single auditable record. "
        "This memo describes a platform that fixes this at the infrastructure level.",
    )

    # ── SECTION 2: What we built ──────────────────────────────────────────────
    section_heading("What We Built")
    body_para(
        "The Carbon Data Trust Platform is a shared governed data infrastructure. "
        "Every dataset that enters the system passes through four layers "
        "before it reaches any report or AI response:",
        bold_parts=["Carbon Data Trust Platform"],
    )

    bullet(
        "Data Quality Rules — each dataset carries automated rules: completeness checks, "
        "value-range validation, cross-field consistency. Rules fire on every data entry. "
        "Violations are scored by severity. A rule cannot be deleted if it has recorded "
        "results — every governance decision is durable.",
        bold_prefix="Data Quality Rules",
    )
    bullet(
        "Data Lineage — every output is traceable end to end: source reading → "
        "emission factor → calculation rule → final figure. "
        "An auditor can open any number and follow it back to the raw meter reading.",
        bold_prefix="Data Lineage",
    )
    bullet(
        "Governed Catalog — every table, field, and dataset is registered with owner, "
        "classification, domain tag, glossary term, and live quality score. "
        "Changes go through a steward approval workflow. Deletions are blocked by policy.",
        bold_prefix="Governed Catalog",
    )
    bullet(
        "Role-Based Access — each department steward manages only their data, "
        "scoped to their organisational unit. Management sees aggregated views. "
        "Access is governed, not informal.",
        bold_prefix="Role-Based Access",
    )

    body_para(
        "This platform layer is domain-agnostic. The Carbon Footprint app is the first "
        "domain running on it — live today: 2,958 t CO₂e across Smart Village Campus "
        "(2023–present), Scope 1, 2, and 3, fully traced and audited."
    )

    p_pulse = doc.add_paragraph()
    set_para_spacing(p_pulse, before=8, after=2)
    rp1 = p_pulse.add_run("Pulse — AI Workspace")
    rp1.bold = True
    rp1.font.color.rgb = ACCENT
    rp1.font.name = "Calibri"
    rp1.font.size = Pt(10.5)
    body_para(
        "Pulse is the AI layer built into the platform. Unlike a generic assistant, "
        "it works directly against your governed, quality-scored records — not a summary of them. "
        "It operates across all domain apps from one workspace."
    )
    bullet(
        "Natural language queries — ask in Arabic or English and get answers traced to source: "
        '"What drove electricity costs up this summer?" returns the specific building '
        "readings and the calculation rule that produced the figure.",
        bold_prefix="Natural language queries",
    )
    bullet(
        "DQ rule suggestions — after profiling a dataset, Pulse suggests quality rules with rationale. "
        "A steward approves or rejects each one. Approved rules become permanent governance objects.",
        bold_prefix="DQ rule suggestions",
    )
    bullet(
        "Report drafting — compliance and sustainability reports grounded in live audited numbers, "
        "GHG Protocol-aligned, with calculation methodology embedded.",
        bold_prefix="Report drafting",
    )
    bullet(
        "Anomaly detection — flags unusual patterns in incoming data before they reach any report. "
        "Surfaced as reviewable items, not silent failures.",
        bold_prefix="Anomaly detection",
    )
    bullet(
        "Cross-domain — the same Pulse workspace handles carbon queries, academic KPI analysis, "
        "and any future domain app. One AI coworker, one governed context.",
        bold_prefix="Cross-domain",
    )

    # ── SECTION 3: Domain apps table ─────────────────────────────────────────
    section_heading("Domain Applications — Current and Planned")
    body_para(
        "Each app runs on the shared platform. Adding a new one does not require rebuilding "
        "the infrastructure — only building the domain logic on top of it.",
    )

    doc.add_paragraph()
    apps_data = [
        ("Application", "Status", "What it does"),
        ("Carbon Footprint Tracking", "Live", "GHG Protocol-aligned emissions across all scopes and campuses"),
        ("Performarc — Academic KPIs", "In migration",
         "Staff and faculty workload, departmental KPIs, publications, grants, evaluation cycles — all branches and colleges"),
        ("Sustainability Goals Tracker", "Proposed",
         "Track AASTMT commitments against SDGs and Egypt Vision 2030: targets, progress, gaps per campus"),
        ("Facilities Management & Lifecycle", "Proposed",
         "Campus assets, maintenance schedules, equipment lifecycle, energy per building"),
        ("Research Projects & Publications", "Proposed",
         "Grant lifecycle, publication tracking, AI coworker for drafting and deadline management"),
    ]

    tbl = doc.add_table(rows=len(apps_data), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Cm(5.5), Cm(3.0), Cm(8.5)]
    STATUS_COLORS = {
        "Live": ("DCFCE7", "15803D"),
        "In migration": ("FEF9C3", "854D0E"),
        "Proposed": ("F1F5F9", "475569"),
    }

    for r_idx, row_data in enumerate(apps_data):
        row = tbl.rows[r_idx]
        for c_idx, (cell_text, width) in enumerate(zip(row_data, col_widths)):
            row.cells[c_idx].width = width
            p2 = row.cells[c_idx].paragraphs[0]
            if r_idx == 0:
                set_cell_bg(row.cells[c_idx], HEADER_BG.replace("#", ""))
                rr2 = p2.add_run(cell_text)
                rr2.bold = True
                rr2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                rr2.font.name = "Calibri"
                rr2.font.size = Pt(9)
            else:
                rr2 = p2.add_run(cell_text)
                rr2.font.name = "Calibri"
                rr2.font.size = Pt(9.5)
                if r_idx % 2 == 0:
                    set_cell_bg(row.cells[c_idx], "F8FBFF")
                if c_idx == 1 and cell_text in STATUS_COLORS:
                    bg, fg = STATUS_COLORS[cell_text]
                    set_cell_bg(row.cells[c_idx], bg)
                    rr2.bold = True
                    rr2.font.color.rgb = RGBColor(
                        int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
                    )
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after = Pt(2)

    set_cell_borders(tbl)
    doc.add_paragraph()

    # ── SECTION 4: What I need ────────────────────────────────────────────────
    section_heading("What I Need From You")
    body_para("These three items are the only blockers. The technical work moves immediately once they are cleared.")

    numbered(
        1,
        "Formal mandate: A directive from your office that this platform is the authorised "
        "data layer for these domains. Without it, each department continues working around it.",
        bold_prefix="Formal mandate:",
    )
    numbered(
        2,
        "Integration access: Authorisation to connect the platform to the ERP and building "
        "management systems. This step is waiting on IT committee sign-off.",
        bold_prefix="Integration access:",
    )
    numbered(
        3,
        "Scope approval: Proceed with the Performarc migration and begin scoping the "
        "Sustainability Goals Tracker and Facilities app.",
        bold_prefix="Scope approval:",
    )

    # ── footer rule ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    set_para_spacing(fp, before=8, after=0)
    fr = fp.add_run("Confidential — for internal circulation")
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED
    fr.font.name = "Calibri"
    fr.italic = True
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    path = os.path.join(OUTPUT_DIR, "chairman_memo_EN.docx")
    doc.save(path)
    print(f"Saved: {path}")



# ── ARABIC ───────────────────────────────────────────────────────────────────

def _ar_run(run):
    """Mark a run as RTL text."""
    rPr = run._r.get_or_add_rPr()
    rPr.append(OxmlElement("w:rtl"))


def _ar_p(p):
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _tbl_bidi(tbl):
    """Let Word render table columns right-to-left without any column flipping."""
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    bv = OxmlElement("w:bidiVisual")
    tblPr.append(bv)


def build_arabic():
    doc = Document()
    page_setup(doc)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    # document-level paragraph bidi default
    doc.styles["Normal"].element.get_or_add_pPr().append(OxmlElement("w:bidi"))

    # ── helpers ───────────────────────────────────────────────────────────────

    def ap(text, bold=False, size=11, color=None, before=0, after=5):
        p = doc.add_paragraph()
        set_para_spacing(p, before=before, after=after)
        _ar_p(p)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Arial"
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        _ar_run(r)
        return p

    def ar_section(text):
        p = doc.add_paragraph()
        set_para_spacing(p, before=10, after=3)
        _ar_p(p)
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = ACCENT
        _ar_run(r)
        pPr2 = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "005CA8")
        pBdr.append(bot)
        pPr2.append(pBdr)

    def ar_bullet(text):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=3)
        _ar_p(p)
        p.paragraph_format.right_indent = Cm(0.5)
        r = p.add_run("• " + text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        _ar_run(r)

    def ar_numbered(num, label, rest):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=4)
        _ar_p(p)
        p.paragraph_format.right_indent = Cm(0.5)
        for txt, bold, color in [
            (num + ".  ", True, ACCENT),
            (label + "  ", True, None),
            (rest, False, None),
        ]:
            r = p.add_run(txt)
            r.bold = bold
            r.font.name = "Arial"
            r.font.size = Pt(11)
            if color:
                r.font.color.rgb = color
            _ar_run(r)

    def ar_cell(cell, text, bold=False, white=False, bg=None, fg=None, size=10):
        p = cell.paragraphs[0]
        _ar_p(p)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Arial"
        r.font.size = Pt(size)
        if white:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif fg:
            r.font.color.rgb = RGBColor(
                int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
            )
        if bg:
            set_cell_bg(cell, bg)
        _ar_run(r)

    # ── title ─────────────────────────────────────────────────────────────────
    ap("مـذكـرة داخلية", bold=True, size=20, color=ACCENT, after=2)
    ap("الأكاديمية العربية للعلوم والتكنولوجيا والنقل البحري",
       size=9, color=MUTED, after=10)

    # ── header table ──────────────────────────────────────────────────────────
    # col 0 = label (shaded, narrow), col 1 = value — bidiVisual flips display order
    htbl = doc.add_table(rows=5, cols=2)
    htbl.style = "Table Grid"
    _tbl_bidi(htbl)
    header_rows = [
        ("إلى",      "سيادة الأستاذ الدكتور / [اسم الرئيس]، رئيس مجلس الأمناء"),
        ("من",       "[اسمك]، مكتب الذكاء الرقمي"),
        ("التاريخ",  "٢٦ أغسطس ٢٠٢٦"),
        ("الموضوع",  "منصة البيانات المؤسسية — ما تم إنجازه وما هو قادم"),
        ("",         "تحية طيبة وبعد،"),
    ]
    for i, (lbl, val) in enumerate(header_rows):
        row = htbl.rows[i]
        row.cells[0].width = Cm(2.5)
        ar_cell(row.cells[0], lbl, bold=True, bg=LIGHT_BG, fg="005CA8", size=9)
        ar_cell(row.cells[1], val, bold=(lbl == "الموضوع"), size=9)
    set_cell_borders(htbl)
    doc.add_paragraph()

    # ── SECTION 1 ─────────────────────────────────────────────────────────────
    ar_section("الوضع الحالي")
    ap(
        "تعمل الأكاديمية حالياً بأنظمة مؤسسية منفصلة: التقارير البيئية، وإدارة أداء "
        "أعضاء هيئة التدريس والعاملين، والمرافق، والبحث العلمي — كلٌّ منها في جزيرة "
        "معلوماتية مستقلة. إعداد أي تقرير يجمع بيانات هذه الأنظمة يستغرق أسابيع، "
        "وجودة البيانات غير موثوقة، ولا يوجد سجل تدقيق موحد. "
        "هذه المذكرة تعرض منصةً تعالج هذا الإشكال من أساسه."
    )

    # ── SECTION 2 ─────────────────────────────────────────────────────────────
    ar_section("ما تم إنجازه")
    ap(
        "قمنا ببناء Data Trust Platform (منصة بيانات كربون الموثوقة) — بنية تحتية محكومة تمر فيها كل بيانة "
        "عبر أربع طبقات قبل وصولها إلى أي تقرير أو استجابة من الذكاء الاصطناعي:"
    )

    def ar_bullet_bold(label, rest):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=3)
        _ar_p(p)
        p.paragraph_format.right_indent = Cm(0.5)
        for txt, bold in [("• " + label + " — ", True), (rest, False)]:
            r = p.add_run(txt)
            r.bold = bold
            r.font.name = "Arial"
            r.font.size = Pt(11)
            _ar_run(r)

    ar_bullet_bold(
        "Data Quality Rules (قواعد جودة البيانات)",
        "لكل مجموعة بيانات قواعد تلقائية: فحوصات الاكتمال، والتحقق من نطاقات القيم، "
        "واتساق الحقول. تُطبَّق القواعد عند كل إدخال، وتُصنَّف المخالفات حسب الخطورة. "
        "كل جدول يحمل درجة جودة حية، ولا يمكن حذف أي قاعدة إذا كانت لها نتائج مسجَّلة.",
    )
    ar_bullet_bold(
        "Data Lineage (أثر البيانات)",
        "كل رقم يمكن تتبعه من أوله إلى آخره: قراءة المصدر ← عامل الانبعاث ← "
        "قاعدة الحساب ← المخرج النهائي. "
        "يمكن لأي مدقق فتح أي مخرج والوصول إلى القراءة الخام التي أنتجته.",
    )
    ar_bullet_bold(
        "Data Catalog (الكتالوج الرقمي)",
        "كل جدول وحقل ومجموعة بيانات مُسجَّلة مع المسؤول عنها، والتصنيف، "
        "والمصطلح المعياري، ودرجة الجودة الحية. "
        "التعديلات تمر بسير عمل اعتماد، والحذف محكوم بسياسات قابلة للضبط.",
    )
    ar_bullet_bold(
        "Access Control (التحكم في الوصول)",
        "كل مسؤول قطاع يرى ويدير بياناته فقط بحسب وحدته التنظيمية. "
        "الإدارة العليا ترى الصورة الموحدة. الصلاحيات محكومة، لا اتفاقية.",
    )

    ap(
        "هذه الطبقة محايدة من حيث المجال. تطبيق البصمة الكربونية هو أول تطبيق يعمل عليها، "
        "وهو نشط حالياً: ٢٩٥٨ طن مكافئ CO₂ لحرم القرية الذكية (٢٠٢٣ حتى الآن)، "
        "النطاقات ١ و٢ و٣، بأثر تدقيقي كامل.",
        before=5,
    )

    p_pulse = doc.add_paragraph()
    set_para_spacing(p_pulse, before=8, after=2)
    _ar_p(p_pulse)
    rph = p_pulse.add_run("Pulse — AI Workspace (مساحة الذكاء الاصطناعي)")
    rph.bold = True
    rph.font.color.rgb = ACCENT
    rph.font.name = "Arial"
    rph.font.size = Pt(11)
    _ar_run(rph)
    ap(
        "Pulse هو طبقة الذكاء الاصطناعي المدمجة في المنصة. خلافاً للمساعدات الذكية العامة، "
        "يصل Pulse إلى سجلاتك الرسمية المحكومة — لا إلى نسخة ملخصة منها. يعمل عبر جميع تطبيقات المنصة من مساحة عمل واحدة.",
        after=3,
    )
    ar_bullet_bold(
        "Natural Language Queries (استعلامات باللغة الطبيعية)",
        "اطرح سؤالك بالعربية أو الإنجليزية واحصل على إجابة مربوطة بالسجلات المصدرية: "
        '"ما سبب ارتفاع فاتورة الكهرباء؟" '
        "تُعيدك إلى قراءات المبنى المحدد وقاعدة الحساب التي أنتجت الرقم.",
    )
    ar_bullet_bold(
        "DQ Rule Suggestions (اقتراحات قواعد الجودة)",
        "بعد تحليل مجموعة بيانات، يقترح Pulse قواعد جودة مع مبرراتها. "
        "المسؤول يعتمدها أو يرفضها. القواعد المعتمدة تصبح كائنات حوكمة دائمة.",
    )
    ar_bullet_bold(
        "Report Drafting (صياغة التقارير)",
        "تقارير الامتثال والاستدامة مبنية على أرقام حية مدققة، متوافقة مع بروتوكول GHG، "
        "مع منهجية الحساب مضمَّنة في متن التقرير.",
    )
    ar_bullet_bold(
        "Anomaly Detection (كشف الشذوذات)",
        "يرصد Pulse الأنماط غير الطبيعية في البيانات الواردة قبل وصولها إلى أي تقرير. "
        "التنبيهات تظهر كبنود مراجعة مع مبرراتها، لا كإخفاقات صامتة.",
    )
    ar_bullet_bold(
        "Cross-Domain (عبر جميع التطبيقات)",
        "نفس مساحة Pulse تخدم استفسارات الكربون وتحليل مؤشرات الأداء الأكاديمي وأي تطبيق قادم — "
        "مساعد ذكاء اصطناعي واحد، سياق محكوم واحد.",
    )

    # ── SECTION 3: apps table ─────────────────────────────────────────────────
    ar_section("تطبيقات المجالات — القائمة والمقترحة")
    ap(
        "المنصة مصممة لاستضافة تطبيقات متخصصة متعددة. كل تطبيق يعمل على نفس البنية "
        "التحتية المشتركة، وإضافة تطبيق جديد لا تستلزم إعادة بناء أي مكوّن."
    )

    doc.add_paragraph()
    apps_ar = [
        ("التطبيق",                                "الحالة",        "ماذا يفعل"),
        ("تتبع البصمة الكربونية",                  "نشط",
         "قياس انبعاثات الغازات الدفيئة عبر جميع النطاقات والحرم الجامعية"),
        ("Performarc — مؤشرات الأداء الأكاديمي",  "قيد الترحيل",
         "عبء العمل للموظفين وأعضاء هيئة التدريس، مؤشرات الأداء على مستوى الأقسام، "
         "المنشورات، المنح البحثية، دورات التقييم — جميع الفروع والكليات"),
        ("متتبع أهداف الاستدامة",                  "مقترح",
         "رصد التزامات الأكاديمية تجاه أهداف التنمية المستدامة ورؤية ٢٠٣٠: "
         "الأهداف والتقدم المحرز والفجوات لكل حرم"),
        ("إدارة المرافق والبنية التحتية",           "مقترح",
         "أصول الحرم، جداول الصيانة، دورة حياة المعدات، استهلاك الطاقة لكل مبنى"),
        ("إدارة البحث العلمي والمنشورات",           "مقترح",
         "دورة حياة المنح، تتبع المنشورات، مساعد ذكاء اصطناعي "
         "لصياغة التقارير وإدارة المواعيد"),
    ]
    STATUS_AR = {
        "نشط":          ("DCFCE7", "15803D"),
        "قيد الترحيل":  ("FEF9C3", "854D0E"),
        "مقترح":        ("F1F5F9", "475569"),
    }
    # col 0 = app name, col 1 = status, col 2 = description
    # bidiVisual makes Word display them right-to-left automatically
    tbl = doc.add_table(rows=len(apps_ar), cols=3)
    _tbl_bidi(tbl)
    col_widths = [Cm(5.0), Cm(2.8), Cm(9.2)]
    for r_idx, (app, status, desc) in enumerate(apps_ar):
        row = tbl.rows[r_idx]
        for c, w in enumerate(col_widths):
            row.cells[c].width = w
        if r_idx == 0:
            for c, txt in enumerate([app, status, desc]):
                ar_cell(row.cells[c], txt, bold=True, white=True, bg="005CA8", size=9)
        else:
            bg_row = "F8FBFF" if r_idx % 2 == 0 else None
            st_bg, st_fg = STATUS_AR.get(status, ("F1F5F9", "475569"))
            ar_cell(row.cells[0], app,    size=10, bg=bg_row)
            ar_cell(row.cells[1], status, size=9,  bg=st_bg, fg=st_fg, bold=True)
            ar_cell(row.cells[2], desc,   size=9,  bg=bg_row)
    set_cell_borders(tbl)
    doc.add_paragraph()

    # ── SECTION 4 ─────────────────────────────────────────────────────────────
    ar_section("ما نحتاجه من سيادتكم")
    ap(
        "هذه النقاط الثلاث هي العوائق الوحيدة أمام التنفيذ. "
        "العمل التقني جاهز للانطلاق فور إقرارها."
    )
    ar_numbered(
        "١", "تفويض رسمي:",
        "توجيه من مكتب سيادتكم يُقرّ هذه المنصة باعتبارها طبقة البيانات المعتمدة "
        "لهذه المجالات. بدون هذا التفويض ستواصل كل إدارة العمل بأسلوبها المستقل."
    )
    ar_numbered(
        "٢", "إذن التكامل:",
        "تفويض بربط المنصة بنظام ERP ونظام إدارة المباني. "
        "هذه الخطوة موقوفة حالياً على موافقة لجنة تقنية المعلومات."
    )
    ar_numbered(
        "٣", "الموافقة على التوسع:",
        "الموافقة على المضي في ترحيل Performarc، والبدء في تحديد نطاق "
        "تطبيقَي متتبع الاستدامة وإدارة المرافق."
    )

    doc.add_paragraph()
    ap("وتفضلوا بقبول فائق الاحترام والتقدير،", size=10, color=MUTED, after=2)
    ap("[اسمك] — مكتب الذكاء الرقمي", size=10, after=0)

    doc.add_paragraph()
    fp = doc.add_paragraph()
    set_para_spacing(fp, before=8, after=0)
    _ar_p(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("سري — للتداول الداخلي")
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED
    fr.font.name = "Arial"
    fr.italic = True
    _ar_run(fr)

    path = os.path.join(OUTPUT_DIR, "chairman_memo_AR.docx")
    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    build_english()
    build_arabic()
    print("Done.")
